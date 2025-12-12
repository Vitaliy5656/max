"""
RAG (Retrieval-Augmented Generation) Engine.

Features:
- Load documents: PDF, DOCX, TXT, MD
- Chunk text with overlap
- Vector index via embeddings
- Semantic search across documents
"""
import uuid
import json
from pathlib import Path
from typing import Optional
from dataclasses import dataclass

import aiosqlite
import tiktoken

# Document parsers
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from .config import config
from .lm_client import lm_client


def _escape_like(query: str) -> str:
    """Escape special characters for SQL LIKE queries (P1 fix: SQL injection)."""
    return query.replace("\\", "\\\\").replace("%", "\\%").replace("_", "\\_")


@dataclass
class Document:
    """Represents an indexed document."""
    id: str
    filename: str
    file_path: str
    file_type: str
    chunk_count: int
    created_at: Optional[str] = None


@dataclass
class Chunk:
    """Document chunk with embedding."""
    id: Optional[int] = None
    document_id: str = ""
    content: str = ""
    chunk_index: int = 0
    tokens: int = 0
    score: float = 0.0  # Relevance score for search results
    source_filename: Optional[str] = None  # Source document filename for context


class RAGEngine:
    """
    Retrieval-Augmented Generation engine.

    Indexes documents and enables semantic search for context augmentation.
    """

    def __init__(self, db: Optional[aiosqlite.Connection] = None):
        self._db = db
        # P2 fix: Use config instead of hardcoded values
        self._chunk_size = config.rag.chunk_size
        self._chunk_overlap = config.rag.chunk_overlap
        # Use tiktoken for accurate token counting
        self._encoder = tiktoken.get_encoding("cl100k_base")

    async def initialize(self, db: aiosqlite.Connection):
        """Initialize with database connection."""
        self._db = db

    def count_tokens(self, text: str) -> int:
        """Count tokens using tiktoken for accuracy."""
        return len(self._encoder.encode(text))
        
    async def add_document(self, file_path: str) -> Document:
        """
        Add and index a document.

        Args:
            file_path: Path to document (PDF, DOCX, TXT, MD)

        Returns:
            Document metadata with chunk count
        """
        path = Path(file_path).expanduser().resolve()

        if not path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")

        # Logic Fix: Deduplication
        # Check if file with same name already exists
        async with self._db.execute(
            "SELECT id FROM documents WHERE filename = ? LIMIT 1",
            (path.name,)
        ) as cursor:
            existing = await cursor.fetchone()
            if existing:
                return await self.get_document(existing["id"])

        # Parse document
        content = await self._parse_document(path)

        if not content.strip():
            raise ValueError("Document is empty or couldn't be parsed")

        # Create document record
        doc_id = str(uuid.uuid4())
        doc = Document(
            id=doc_id,
            filename=path.name,
            file_path=str(path),
            file_type=path.suffix.lower()[1:],
            chunk_count=0
        )

        # Use transaction for atomicity - rollback on any error
        try:
            await self._db.execute(
                """INSERT INTO documents (id, filename, file_path, file_type, chunk_count)
                   VALUES (?, ?, ?, ?, 0)""",
                (doc.id, doc.filename, doc.file_path, doc.file_type)
            )

            # Chunk and index
            chunks = self._split_into_chunks(content)

            for i, chunk_text in enumerate(chunks):
                # Get embedding
                embedding = await lm_client.get_embedding(chunk_text)
                embedding_blob = json.dumps(embedding).encode() if embedding else None

                # Use real token count instead of word count
                token_count = self.count_tokens(chunk_text)

                await self._db.execute(
                    """INSERT INTO document_chunks
                       (document_id, content, embedding, chunk_index, tokens)
                       VALUES (?, ?, ?, ?, ?)""",
                    (doc.id, chunk_text, embedding_blob, i, token_count)
                )

            # Update chunk count
            doc.chunk_count = len(chunks)
            await self._db.execute(
                "UPDATE documents SET chunk_count = ? WHERE id = ?",
                (doc.chunk_count, doc.id)
            )

            await self._db.commit()
            return doc

        except Exception as e:
            # Rollback on error
            await self._db.rollback()
            raise RuntimeError(f"Failed to add document: {e}") from e
    
    async def _parse_document(self, path: Path) -> str:
        """Parse document content based on file type."""
        suffix = path.suffix.lower()
        
        if suffix == ".pdf":
            return self._parse_pdf(path)
        elif suffix == ".docx":
            return self._parse_docx(path)
        elif suffix in (".txt", ".md", ".markdown"):
            return path.read_text(encoding="utf-8", errors="replace")
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
    
    def _parse_pdf(self, path: Path) -> str:
        """Extract text from PDF."""
        if not HAS_PYMUPDF:
            raise ImportError("PyMuPDF not installed. Run: pip install pymupdf")
        
        text_parts = []
        with fitz.open(path) as doc:
            for page in doc:
                text_parts.append(page.get_text())
        
        return "\n".join(text_parts)
    
    def _parse_docx(self, path: Path) -> str:
        """Extract text from DOCX."""
        if not HAS_DOCX:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        doc = DocxDocument(path)
        return "\n".join([p.text for p in doc.paragraphs])
    
    def _split_into_chunks(self, text: str) -> list[str]:
        """Split text into overlapping chunks."""
        words = text.split()
        chunks = []
        
        start = 0
        while start < len(words):
            end = start + self._chunk_size
            chunk_words = words[start:end]
            chunks.append(" ".join(chunk_words))
            
            # Move start with overlap
            start = end - self._chunk_overlap
            if start >= len(words):
                break
        
        return chunks
    
    async def query(
        self,
        question: str,
        top_k: int = 5,
        document_id: Optional[str] = None
    ) -> list[Chunk]:
        """
        Search for relevant chunks.

        Args:
            question: Query to search for
            top_k: Maximum chunks to return
            document_id: Limit search to specific document

        Returns:
            List of relevant chunks sorted by score, with source info
        """
        # Get query embedding
        query_embedding = await lm_client.get_embedding(question)

        if not query_embedding:
            # Fallback to text search if embeddings fail
            return await self._text_search(question, top_k, document_id)

        # Get all chunks with document info (JOIN for source filename)
        if document_id:
            sql = """SELECT c.*, d.filename as source_filename
                     FROM document_chunks c
                     JOIN documents d ON c.document_id = d.id
                     WHERE c.document_id = ?"""
            params = (document_id,)
        else:
            sql = """SELECT c.*, d.filename as source_filename
                     FROM document_chunks c
                     JOIN documents d ON c.document_id = d.id"""
            params = ()

        async with self._db.execute(sql, params) as cursor:
            rows = await cursor.fetchall()

        # Calculate similarity scores
        scored_chunks = []
        for row in rows:
            if row["embedding"]:
                chunk_embedding = json.loads(row["embedding"].decode())
                score = self._cosine_similarity(query_embedding, chunk_embedding)
            else:
                # Text-based fallback score
                score = self._text_similarity(question, row["content"])

            scored_chunks.append(Chunk(
                id=row["id"],
                document_id=row["document_id"],
                content=row["content"],
                chunk_index=row["chunk_index"],
                tokens=row["tokens"],
                score=score,
                source_filename=row["source_filename"]
            ))

        # Sort by score and return top_k
        scored_chunks.sort(key=lambda c: c.score, reverse=True)
        return scored_chunks[:top_k]
    
    async def _text_search(
        self,
        query: str,
        top_k: int,
        document_id: Optional[str]
    ) -> list[Chunk]:
        """Fallback text-based search."""
        # P1 fix: Escape special SQL LIKE characters
        escaped_query = _escape_like(query)
        if document_id:
            sql = """SELECT c.*, d.filename as source_filename
                     FROM document_chunks c
                     JOIN documents d ON c.document_id = d.id
                     WHERE c.document_id = ? AND c.content LIKE ? ESCAPE '\\'
                     LIMIT ?"""
            params = (document_id, f"%{escaped_query}%", top_k)
        else:
            sql = """SELECT c.*, d.filename as source_filename
                     FROM document_chunks c
                     JOIN documents d ON c.document_id = d.id
                     WHERE c.content LIKE ? ESCAPE '\\'
                     LIMIT ?"""
            params = (f"%{escaped_query}%", top_k)

        async with self._db.execute(sql, params) as cursor:
            rows = await cursor.fetchall()

        return [
            Chunk(
                id=row["id"],
                document_id=row["document_id"],
                content=row["content"],
                chunk_index=row["chunk_index"],
                tokens=row["tokens"],
                score=1.0,
                source_filename=row["source_filename"]
            )
            for row in rows
        ]
    
    def _cosine_similarity(self, a: list[float], b: list[float]) -> float:
        """Calculate cosine similarity between two vectors."""
        if len(a) != len(b):
            return 0.0
        
        dot_product = sum(x * y for x, y in zip(a, b))
        norm_a = sum(x * x for x in a) ** 0.5
        norm_b = sum(x * x for x in b) ** 0.5
        
        if norm_a == 0 or norm_b == 0:
            return 0.0
        
        return dot_product / (norm_a * norm_b)
    
    def _text_similarity(self, query: str, text: str) -> float:
        """Simple text-based similarity score."""
        query_words = set(query.lower().split())
        text_words = set(text.lower().split())
        
        if not query_words:
            return 0.0
        
        return len(query_words & text_words) / len(query_words)
    
    async def list_documents(self) -> list[Document]:
        """List all indexed documents."""
        async with self._db.execute(
            "SELECT * FROM documents ORDER BY created_at DESC"
        ) as cursor:
            rows = await cursor.fetchall()
        
        return [
            Document(
                id=row["id"],
                filename=row["filename"],
                file_path=row["file_path"],
                file_type=row["file_type"],
                chunk_count=row["chunk_count"],
                created_at=row["created_at"]
            )
            for row in rows
        ]
    
    async def get_document(self, doc_id: str) -> Optional[Document]:
        """Get document by ID."""
        async with self._db.execute(
            "SELECT * FROM documents WHERE id = ?", (doc_id,)
        ) as cursor:
            row = await cursor.fetchone()
        
        if not row:
            return None
        
        return Document(
            id=row["id"],
            filename=row["filename"],
            file_path=row["file_path"],
            file_type=row["file_type"],
            chunk_count=row["chunk_count"],
            created_at=row["created_at"]
        )
    
    async def remove_document(self, doc_id: str) -> bool:
        """Remove document and its chunks."""
        # Delete chunks first (foreign key)
        await self._db.execute(
            "DELETE FROM document_chunks WHERE document_id = ?",
            (doc_id,)
        )
        
        # Delete document
        cursor = await self._db.execute(
            "DELETE FROM documents WHERE id = ?",
            (doc_id,)
        )
        await self._db.commit()
        
        return cursor.rowcount > 0
    
    async def get_context_for_query(
        self,
        question: str,
        max_tokens: int = 2000
    ) -> str:
        """
        Get formatted context from relevant documents for a query.
        Used to augment LLM prompts with document knowledge.
        """
        chunks = await self.query(question, top_k=10)
        
        if not chunks:
            return ""
        
        context_parts = []
        tokens_used = 0
        
        for chunk in chunks:
            chunk_tokens = chunk.tokens or len(chunk.content.split())
            if tokens_used + chunk_tokens > max_tokens:
                break
            
            context_parts.append(chunk.content)
            tokens_used += chunk_tokens
        
        if not context_parts:
            return ""
        
        return f"""[Релевантная информация из документов:]

{chr(10).join(context_parts)}

[Конец контекста из документов]"""


# Global RAG engine
rag = RAGEngine()
